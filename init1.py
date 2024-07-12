import pandas as pd
from docx import Document
from collections import defaultdict
from datetime import datetime

# 读取CSV文件
def read_log(file_path):
    df = pd.read_csv(file_path)
    return df

# 筛选特定时间范围内的数据
def filter_by_time(df, start_time, end_time):
    df['时间'] = pd.to_datetime(df['时间'])
    filtered_df = df.loc[(df['时间'] >= start_time) & (df['时间'] <= end_time)]
    return filtered_df

# 分析日志数据并打印分组及数量
def analyze_log(df):
    # 按照防护类型分组并统计数量
    df.loc[:, '防护类型'] = df['防护类型'].str.lower()  # 将防护类型转换为小写
    type_summary = df.groupby('防护类型').size().reset_index(name='数量')
    print("防护类型分组及其数量：")
    print(type_summary)

    # 按照防护子类型分组并统计数量
    subtype_summary = df.groupby('防护子类型').size().reset_index(name='数量')
    print("\n防护子类型分组及其数量：")
    print(subtype_summary)
    return type_summary, subtype_summary

# 获取当前日期并格式化
def get_current_date():
    return datetime.now().strftime('%Y年%m月%d日')

# 获取当前月份
def get_current_month():
    return datetime.now().strftime('%Y年%m月')

# 获取当前日期的天数部分
def get_current_day():
    return datetime.now().strftime('%d日')

# 从Excel文件中读取IP封禁列表并计算ipcount
def get_ip_count(excel_file):
    df = pd.read_excel(excel_file, sheet_name='hvv恶意IP封禁')
    ip_count = df.iloc[:, -1].nunique()  # 获取最后一列并计算唯一值的数量
    return ip_count

# 生成日报内容
def generate_daily_report(type_summary, start_time, end_time, ipcount):
    # 从统计数据中提取数字，如果不存在则默认为0
    counts = defaultdict(int)
    for _, row in type_summary.iterrows():
        counts[row['防护类型']] = row['数量']

    total_attacks = type_summary['数量'].sum()
    print(total_attacks)
    # 格式化时间段
    start_time_str = start_time.strftime('%Y年%m月%d日%H:%M').replace(' 0', ' ').replace('-0', '-')
    end_time_str = end_time.strftime('%Y年%m月%d日%H:%M').replace(' 0', ' ').replace('-0', '-')

    report = {
        "开始时间": start_time_str,
        "结束时间": end_time_str,
        "ip数量": f"{ipcount}",
        "ddos次数": counts['ddos攻击'],
        "注入次数": counts['注入攻击'],
        "跨站攻击次数": counts['跨站攻击'],
        "信息泄露次数": counts['信息泄露'],
        "探测访问次数": counts['探测访问'],
        "特殊漏洞攻击次数": counts['特殊漏洞攻击'],
        "资源非法访问次数": counts['资源非法访问'],
        "恶意软件次数": counts['恶意软件'],
        "总攻击次数": total_attacks,
        "本月": get_current_date(),
        "本日": get_current_date()
    }

    return report

# 替换模板中的变量并保存新的Word文档，同时保留格式
def save_report_to_word(template_path, report, file_path):
    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        for key, value in report.items():
            if f"{{{key}}}" in paragraph.text:
                for run in paragraph.runs:
                    if f"{{{key}}}" in run.text:
                        run.text = run.text.replace(f"{{{key}}}", str(value))

    # 替换表格中的变量
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in report.items():
                        if f"{{{key}}}" in paragraph.text:
                            for run in paragraph.runs:
                                if f"{{{key}}}" in run.text:
                                    run.text = run.text.replace(f"{{{key}}}", str(value))

    doc.save(file_path)

# 主函数
def main(log_file, template_path, output_directory, ip_excel_file):
    log_df = read_log(log_file)
    start_time = pd.to_datetime('2024-07-11 18:00:00')
    end_time = pd.to_datetime('2024-07-12 18:00:00')
    filtered_log_df = filter_by_time(log_df, start_time, end_time)
    type_summary, subtype_summary = analyze_log(filtered_log_df)
    ip_count = get_ip_count(ip_excel_file)
    print(ip_count)
    report = generate_daily_report(type_summary, start_time, end_time, ip_count)

    # 获取当前日期并生成文件名
    current_date = get_current_date()
    report_file = f"{output_directory}\\{current_date}人社局护网受攻击情况汇总.docx"

    save_report_to_word(template_path, report, report_file)

# 执行程序

#攻击日志路径
log_file = r'C:\Users\one\Downloads\20240709_180005.csv'
# 模版路径
template_path = r'D:\360MoveData\Users\one\Desktop\日志分析记录\人保局日志分析记录仪\人保局日志\模版.docx'
#导出路径
output_directory = r'D:\360MoveData\Users\one\Desktop\日志分析记录\人保局日志分析记录仪\人保局日志'
#封禁IP 文件路径
ip_excel_file = r'D:\360MoveData\Users\one\Desktop\日志分析记录\人保局日志分析记录仪\人保局日志\人社局安全巡检恶意IP封禁列表.xlsx'
main(log_file, template_path, output_directory, ip_excel_file)

print("日报内容生成完毕，并保存为最新的日期文件名。")
